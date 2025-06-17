import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document

def browse_directory():
    path = filedialog.askdirectory()
    dir_var.set(path)

def run_program():
    name = name_var.get()
    ecu = ecu_var.get()
    directory = dir_var.get()

    if not name or not ecu or not directory:
        messagebox.showerror("Missing Input", "Please fill in all fields.")
        return

    try:
        output_file = generate_doc(name, ecu, directory)
        messagebox.showinfo("Success", f"Document created:\n{output_file}")
        os.startfile(output_file)  # Opens the file (works on Windows)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate document:\n{e}")

def generate_doc(name, ecu, directory):
    # 🔁 Replace this logic with your actual program if needed
    doc = Document()
    doc.add_heading(f"Report for {name}", level=1)
    doc.add_paragraph(f"ECU Name: {ecu}")
    doc.add_paragraph("This is a sample DOC generated via GUI.")
    output_path = os.path.join(directory, f"{name}_{ecu}_report.docx")
    doc.save(output_path)
    return output_path

# GUI layout
root = tk.Tk()
root.title("ECU DOC Generator")
root.geometry("450x200")
root.resizable(False, False)

name_var = tk.StringVar()
ecu_var = tk.StringVar()
dir_var = tk.StringVar()

tk.Label(root, text="Enter Name:").grid(row=0, column=0, padx=10, pady=10, sticky='e')
tk.Entry(root, textvariable=name_var, width=40).grid(row=0, column=1, columnspan=2)

tk.Label(root, text="Enter ECU:").grid(row=1, column=0, padx=10, pady=10, sticky='e')
tk.Entry(root, textvariable=ecu_var, width=40).grid(row=1, column=1, columnspan=2)

tk.Label(root, text="Output Directory:").grid(row=2, column=0, padx=10, pady=10, sticky='e')
tk.Entry(root, textvariable=dir_var, width=30).grid(row=2, column=1)
tk.Button(root, text="Browse", command=browse_directory).grid(row=2, column=2)

tk.Button(root, text="Generate DOC File", command=run_program, bg="#4CAF50", fg="white").grid(row=3, column=1, pady=20)

root.mainloop()
